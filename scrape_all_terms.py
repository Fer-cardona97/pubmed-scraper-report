import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document

def search_entrez(query, max_articles=5):
    base_url = "https://pubmed.ncbi.nlm.nih.gov/"
    search_url = f"{base_url}?term={query}"
    articles = []
    page = 1

    while len(articles) < max_articles:
        print(f"Searching page {page} for term '{query}'...")
        response = requests.get(search_url, params={'page': page})
        soup = BeautifulSoup(response.content, 'html.parser')
        results = soup.find_all('article', class_='full-docsum')

        if not results:
            print(f"No more results for term '{query}'.")
            break

        for result in results:
            title = result.find('a', class_='docsum-title').text.strip()
            link = result.find('a', class_='docsum-title')['href']
            full_link = urljoin(base_url, link)
            articles.append((title, full_link))
        
        page += 1

    return articles[:max_articles]

def filter_articles(articles, search_terms):
    filtered_articles = []
    for title, link in articles:
        response = requests.get(link)
        soup = BeautifulSoup(response.content, 'html.parser')
        article_text = soup.get_text().lower()
        if all(term.lower() in article_text for term in search_terms):
            filtered_articles.append((title, link))
    return filtered_articles

def add_search_terms_to_doc(doc, search_terms):
    doc.add_heading("Search Terms", level=1)
    doc.add_paragraph(", ".join(search_terms))

def add_results_to_doc(doc, term, articles):
    doc.add_heading(f"Combined Search Results for {', '.join(term)}", level=1)
    for title, link in articles:
        doc.add_paragraph(title, style='Heading 2')
        doc.add_paragraph(f"Link: {link}")
    doc.add_page_break()

if __name__ == "__main__":
    search_terms = ["anger", "traffic" ]
    max_articles_per_term = 50
    output_dir = "C:/Users/FernanDcardona/Desktop/Experiment/Web_Scrapping_Results/"
    doc = Document()
    combined_articles = []

    for term in search_terms:
        articles = search_entrez(term, max_articles_per_term)

        if not articles:
            print(f"No articles found for term '{term}'.")
            continue

        combined_articles.extend(articles)
        print(f"Searching page {term} for term '{term}'...")
    
    print("Filtering all results...")
    filtered_articles = filter_articles(combined_articles, search_terms)

    if not filtered_articles:
        print("No articles found for any of the terms.")
    else:
        add_search_terms_to_doc(doc, search_terms)
        add_results_to_doc(doc, search_terms, filtered_articles)

    doc.save(f"{output_dir}PubMed_Search_All_Results.docx")
    print("All search results saved to 'PubMed_Search_All_Results.docx'.")
